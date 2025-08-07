#!/usr/bin/env python3
"""
Nexus Document Ingestion System
Network-efficient document scanning with local caching and vector analysis
"""

import os
import shutil
import hashlib
import json
import time
from pathlib import Path
from datetime import datetime
from typing import Dict, List, Optional
import sqlite3
import chromadb
from watchdog.observers import Observer
from watchdog.events import FileSystemEventHandler

class NexusDocumentCache:
    """Local caching system for network documents"""
    
    def __init__(self, cache_root: str = "./nexus_cache"):
        self.cache_root = Path(cache_root)
        self.cache_root.mkdir(exist_ok=True)
        
        # Cache metadata database
        self.cache_db = self.cache_root / "cache_metadata.db"
        self.init_cache_db()
        
    def init_cache_db(self):
        """Initialize cache tracking database"""
        conn = sqlite3.connect(self.cache_db)
        cursor = conn.cursor()
        
        cursor.execute('''
            CREATE TABLE IF NOT EXISTS cached_files (
                id INTEGER PRIMARY KEY,
                network_path TEXT UNIQUE,
                local_path TEXT,
                file_hash TEXT,
                last_modified REAL,
                cached_timestamp TEXT,
                file_size INTEGER,
                status TEXT DEFAULT 'cached'
            )
        ''')
        
        conn.commit()
        conn.close()
    
    def get_file_hash(self, file_path: Path) -> str:
        """Generate hash for file content"""
        hasher = hashlib.md5()
        try:
            with open(file_path, 'rb') as f:
                for chunk in iter(lambda: f.read(4096), b""):
                    hasher.update(chunk)
            return hasher.hexdigest()
        except Exception:
            return ""
    
    def should_update_cache(self, network_path: Path) -> bool:
        """Check if file needs to be cached or updated"""
        if not network_path.exists():
            return False
            
        conn = sqlite3.connect(self.cache_db)
        cursor = conn.cursor()
        
        cursor.execute(
            'SELECT last_modified, file_hash FROM cached_files WHERE network_path = ?',
            (str(network_path),)
        )
        
        result = cursor.fetchone()
        conn.close()
        
        if not result:
            return True  # Not cached yet
            
        cached_mtime, cached_hash = result
        current_mtime = network_path.stat().st_mtime
        
        # Check if file was modified
        if current_mtime > cached_mtime:
            return True
            
        return False
    
    def cache_file(self, network_path: Path) -> Optional[Path]:
        """Cache a network file locally"""
        if not self.should_update_cache(network_path):
            # Return existing cached path
            conn = sqlite3.connect(self.cache_db)
            cursor = conn.cursor()
            cursor.execute(
                'SELECT local_path FROM cached_files WHERE network_path = ?',
                (str(network_path),)
            )
            result = cursor.fetchone()
            conn.close()
            
            if result:
                return Path(result[0])
        
        try:
            # Create local cache structure
            relative_path = network_path.relative_to(network_path.anchor)
            local_path = self.cache_root / "files" / relative_path
            local_path.parent.mkdir(parents=True, exist_ok=True)
            
            # Copy file to cache
            shutil.copy2(network_path, local_path)
            
            # Get file metadata
            file_stat = network_path.stat()
            file_hash = self.get_file_hash(local_path)
            
            # Update cache database
            conn = sqlite3.connect(self.cache_db)
            cursor = conn.cursor()
            
            cursor.execute('''
                INSERT OR REPLACE INTO cached_files 
                (network_path, local_path, file_hash, last_modified, cached_timestamp, file_size)
                VALUES (?, ?, ?, ?, ?, ?)
            ''', (
                str(network_path),
                str(local_path),
                file_hash,
                file_stat.st_mtime,
                datetime.now().isoformat(),
                file_stat.st_size
            ))
            
            conn.commit()
            conn.close()
            
            print(f"✅ Cached: {network_path.name}")
            return local_path
            
        except Exception as e:
            print(f"❌ Cache failed for {network_path}: {e}")
            return None

class NexusDocumentProcessor:
    """Document content extraction and processing"""
    
    def __init__(self):
        self.supported_formats = {
            '.txt': self.extract_text,
            '.md': self.extract_text,
            '.py': self.extract_text,
            '.js': self.extract_text,
            '.html': self.extract_text,
            '.css': self.extract_text,
            '.pdf': self.extract_pdf,
            '.docx': self.extract_docx,
            '.xlsx': self.extract_xlsx,
        }
    
    def extract_text(self, file_path: Path) -> Dict:
        """Extract content from text files"""
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                content = f.read()
            
            return {
                'content': content,
                'word_count': len(content.split()),
                'char_count': len(content),
                'type': 'text'
            }
        except Exception as e:
            return {'error': str(e), 'type': 'text'}
    
    def extract_pdf(self, file_path: Path) -> Dict:
        """Extract content from PDF files"""
        try:
            import PyPDF2
            with open(file_path, 'rb') as f:
                reader = PyPDF2.PdfReader(f)
                content = ""
                for page in reader.pages:
                    content += page.extract_text() + "\n"
            
            return {
                'content': content,
                'page_count': len(reader.pages),
                'word_count': len(content.split()),
                'type': 'pdf'
            }
        except ImportError:
            return {'error': 'PyPDF2 not installed', 'type': 'pdf'}
        except Exception as e:
            return {'error': str(e), 'type': 'pdf'}
    
    def extract_docx(self, file_path: Path) -> Dict:
        """Extract content from Word documents"""
        try:
            from docx import Document
            doc = Document(file_path)
            content = "\n".join([paragraph.text for paragraph in doc.paragraphs])
            
            return {
                'content': content,
                'paragraph_count': len(doc.paragraphs),
                'word_count': len(content.split()),
                'type': 'docx'
            }
        except ImportError:
            return {'error': 'python-docx not installed', 'type': 'docx'}
        except Exception as e:
            return {'error': str(e), 'type': 'docx'}
    
    def extract_xlsx(self, file_path: Path) -> Dict:
        """Extract content from Excel files"""
        try:
            import pandas as pd
            df = pd.read_excel(file_path, sheet_name=None)
            
            content = ""
            for sheet_name, sheet_data in df.items():
                content += f"Sheet: {sheet_name}\n"
                content += sheet_data.to_string(index=False) + "\n\n"
            
            return {
                'content': content,
                'sheet_count': len(df),
                'type': 'xlsx'
            }
        except ImportError:
            return {'error': 'pandas not installed', 'type': 'xlsx'}
        except Exception as e:
            return {'error': str(e), 'type': 'xlsx'}
    
    def process_file(self, file_path: Path) -> Optional[Dict]:
        """Process a cached file and extract content"""
        suffix = file_path.suffix.lower()
        
        if suffix in self.supported_formats:
            extractor = self.supported_formats[suffix]
            result = extractor(file_path)
            
            # Add file metadata
            result.update({
                'file_path': str(file_path),
                'file_name': file_path.name,
                'file_size': file_path.stat().st_size,
                'processed_timestamp': datetime.now().isoformat()
            })
            
            return result
        
        return None

class NexusVectorStore:
    """Vector database integration for Nexus documents"""
    
    def __init__(self, collection_name: str = "nexus_documents"):
        self.client = chromadb.PersistentClient(path="./nexus_vector_db")
        self.collection = self.client.get_or_create_collection(
            name=collection_name,
            metadata={"description": "Nexus document vector storage"}
        )
    
    def add_document(self, doc_data: Dict, file_path: str):
        """Add document to vector database"""
        if 'content' not in doc_data or 'error' in doc_data:
            return False
        
        content = doc_data['content']
        if not content.strip():
            return False
        
        # Create document ID from file path
        doc_id = hashlib.md5(file_path.encode()).hexdigest()
        
        # Create metadata
        metadata = {
            'file_path': file_path,
            'file_name': doc_data.get('file_name', ''),
            'file_type': doc_data.get('type', 'unknown'),
            'word_count': doc_data.get('word_count', 0),
            'processed_timestamp': doc_data.get('processed_timestamp', ''),
        }
        
        try:
            # Add to ChromaDB
            self.collection.add(
                documents=[content],
                metadatas=[metadata],
                ids=[doc_id]
            )
            return True
        except Exception as e:
            print(f"❌ Vector store error: {e}")
            return False
    
    def search_documents(self, query: str, n_results: int = 10) -> List[Dict]:
        """Search documents by semantic similarity"""
        try:
            results = self.collection.query(
                query_texts=[query],
                n_results=n_results
            )
            
            documents = []
            for i in range(len(results['documents'][0])):
                documents.append({
                    'content': results['documents'][0][i],
                    'metadata': results['metadatas'][0][i],
                    'distance': results['distances'][0][i] if 'distances' in results else None
                })
            
            return documents
        except Exception as e:
            print(f"❌ Search error: {e}")
            return []

class NexusDocumentIngestion:
    """Main Nexus document ingestion system"""
    
    def __init__(self, network_paths: List[str]):
        self.network_paths = [Path(p) for p in network_paths]
        self.cache = NexusDocumentCache()
        self.processor = NexusDocumentProcessor()
        self.vector_store = NexusVectorStore()
        
        # Statistics
        self.stats = {
            'files_scanned': 0,
            'files_cached': 0,
            'files_processed': 0,
            'files_vectorized': 0,
            'errors': 0
        }
    
    def scan_network_paths(self):
        """Scan network paths and cache new/updated files"""
        print("🔍 Scanning network paths for documents...")
        
        for network_path in self.network_paths:
            if not network_path.exists():
                print(f"⚠️  Path not accessible: {network_path}")
                continue
            
            print(f"📁 Scanning: {network_path}")
            self.scan_directory(network_path)
    
    def scan_directory(self, directory: Path):
        """Recursively scan directory for documents"""
        try:
            for item in directory.rglob("*"):
                if item.is_file():
                    self.stats['files_scanned'] += 1
                    
                    # Cache file if needed
                    cached_path = self.cache.cache_file(item)
                    if cached_path:
                        self.stats['files_cached'] += 1
                        
                        # Process and vectorize
                        self.process_cached_file(cached_path, str(item))
                    
                    # Progress indicator
                    if self.stats['files_scanned'] % 100 == 0:
                        print(f"   Processed {self.stats['files_scanned']} files...")
                        
        except Exception as e:
            print(f"❌ Directory scan error: {e}")
            self.stats['errors'] += 1
    
    def process_cached_file(self, cached_path: Path, original_path: str):
        """Process a cached file and add to vector store"""
        doc_data = self.processor.process_file(cached_path)
        
        if doc_data:
            self.stats['files_processed'] += 1
            
            # Add to vector database
            if self.vector_store.add_document(doc_data, original_path):
                self.stats['files_vectorized'] += 1
        else:
            # Unsupported file type (not an error)
            pass
    
    def search_documents(self, query: str, limit: int = 10) -> List[Dict]:
        """Search the Nexus document collection"""
        return self.vector_store.search_documents(query, limit)
    
    def get_stats(self) -> Dict:
        """Get ingestion statistics"""
        return self.stats.copy()

def main():
    """Main entry point for Nexus Document Ingestion"""
    print("🔗 Nexus Document Ingestion System")
    print("=" * 40)
    
    # Example network paths (customize for your environment)
    network_paths = [
        "/mnt/shared_drives/documents",
        "/mnt/shared_drives/projects", 
        "/home/jonclaude/Documents",  # Local documents too
    ]
    
    # Filter to existing paths
    existing_paths = []
    for path in network_paths:
        if Path(path).exists():
            existing_paths.append(path)
        else:
            print(f"⚠️  Path not found: {path}")
    
    if not existing_paths:
        print("❌ No accessible paths found. Please update network_paths.")
        return
    
    # Initialize ingestion system
    nexus = NexusDocumentIngestion(existing_paths)
    
    # Start ingestion
    start_time = time.time()
    nexus.scan_network_paths()
    end_time = time.time()
    
    # Show results
    stats = nexus.get_stats()
    print(f"\n📊 Ingestion Complete ({end_time - start_time:.1f}s)")
    print(f"   Files scanned: {stats['files_scanned']}")
    print(f"   Files cached: {stats['files_cached']}")
    print(f"   Files processed: {stats['files_processed']}")
    print(f"   Files vectorized: {stats['files_vectorized']}")
    print(f"   Errors: {stats['errors']}")
    
    # Test search
    if stats['files_vectorized'] > 0:
        print(f"\n🔍 Testing document search...")
        results = nexus.search_documents("project management", limit=3)
        
        for i, result in enumerate(results, 1):
            metadata = result['metadata']
            print(f"{i}. {metadata['file_name']} ({metadata['file_type']})")
            print(f"   Path: {metadata['file_path']}")
            print(f"   Preview: {result['content'][:100]}...")
            print()

if __name__ == "__main__":
    main()