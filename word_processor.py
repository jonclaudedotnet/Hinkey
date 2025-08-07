#!/usr/bin/env python3
"""
Word Document Processor
Extracts text and metadata from Word documents
"""

import os
import sqlite3
from pathlib import Path
from datetime import datetime
import re
import hashlib
from typing import Dict, List, Optional
import json
from docx import Document
import zipfile
import xml.etree.ElementTree as ET

class WordProcessor:
    """Process Word documents and extract structured data"""
    
    def __init__(self, db_path: str = "word_data.db"):
        self.db_path = db_path
        self.conn = sqlite3.connect(db_path)
        self.init_database()
        self.stats = {
            'files_processed': 0,
            'text_extracted': 0,
            'metadata_extracted': 0,
            'tables_found': 0,
            'images_found': 0,
            'errors': []
        }
        
    def init_database(self):
        """Initialize Word document database schema"""
        cursor = self.conn.cursor()
        
        # Main documents table
        cursor.execute('''
            CREATE TABLE IF NOT EXISTS word_docs (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                file_path TEXT UNIQUE,
                file_name TEXT,
                file_size INTEGER,
                file_hash TEXT,
                creation_date TEXT,
                modification_date TEXT,
                author TEXT,
                last_modified_by TEXT,
                title TEXT,
                subject TEXT,
                word_count INTEGER,
                page_count INTEGER,
                paragraph_count INTEGER,
                table_count INTEGER,
                image_count INTEGER,
                document_type TEXT,
                extracted_text TEXT,
                extracted_tables TEXT,
                key_entities TEXT,
                processed_timestamp TIMESTAMP DEFAULT CURRENT_TIMESTAMP
            )
        ''')
        
        # Document sections table
        cursor.execute('''
            CREATE TABLE IF NOT EXISTS word_sections (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                doc_id INTEGER,
                section_type TEXT,
                section_title TEXT,
                section_content TEXT,
                section_order INTEGER,
                FOREIGN KEY (doc_id) REFERENCES word_docs(id)
            )
        ''')
        
        # Extracted entities
        cursor.execute('''
            CREATE TABLE IF NOT EXISTS word_entities (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                doc_id INTEGER,
                entity_type TEXT,
                entity_value TEXT,
                context TEXT,
                FOREIGN KEY (doc_id) REFERENCES word_docs(id)
            )
        ''')
        
        # Create indexes
        cursor.execute('CREATE INDEX IF NOT EXISTS idx_doc_type ON word_docs(document_type)')
        cursor.execute('CREATE INDEX IF NOT EXISTS idx_author ON word_docs(author)')
        
        self.conn.commit()
        
    def get_file_hash(self, file_path: Path) -> str:
        """Calculate SHA256 hash of file"""
        sha256_hash = hashlib.sha256()
        with open(file_path, "rb") as f:
            for byte_block in iter(lambda: f.read(4096), b""):
                sha256_hash.update(byte_block)
        return sha256_hash.hexdigest()
        
    def extract_docx_metadata(self, file_path: Path) -> Dict:
        """Extract metadata from .docx files"""
        metadata = {
            'author': None,
            'last_modified_by': None,
            'created': None,
            'modified': None,
            'title': None,
            'subject': None
        }
        
        try:
            # .docx files are zip archives
            with zipfile.ZipFile(file_path, 'r') as docx:
                # Extract core properties
                if 'docProps/core.xml' in docx.namelist():
                    core_xml = docx.read('docProps/core.xml')
                    root = ET.fromstring(core_xml)
                    
                    # Define namespaces
                    ns = {
                        'cp': 'http://schemas.openxmlformats.org/package/2006/metadata/core-properties',
                        'dc': 'http://purl.org/dc/elements/1.1/',
                        'dcterms': 'http://purl.org/dc/terms/'
                    }
                    
                    # Extract properties
                    creator = root.find('.//dc:creator', ns)
                    if creator is not None and creator.text:
                        metadata['author'] = creator.text
                        
                    last_modifier = root.find('.//cp:lastModifiedBy', ns)
                    if last_modifier is not None and last_modifier.text:
                        metadata['last_modified_by'] = last_modifier.text
                        
                    created = root.find('.//dcterms:created', ns)
                    if created is not None and created.text:
                        metadata['created'] = created.text
                        
                    modified = root.find('.//dcterms:modified', ns)
                    if modified is not None and modified.text:
                        metadata['modified'] = modified.text
                        
                    title = root.find('.//dc:title', ns)
                    if title is not None and title.text:
                        metadata['title'] = title.text
                        
                    subject = root.find('.//dc:subject', ns)
                    if subject is not None and subject.text:
                        metadata['subject'] = subject.text
                        
        except Exception as e:
            pass
            
        return metadata
        
    def extract_text_and_structure(self, file_path: Path) -> Dict:
        """Extract text and document structure"""
        result = {
            'text': '',
            'paragraphs': [],
            'tables': [],
            'sections': [],
            'word_count': 0,
            'paragraph_count': 0,
            'table_count': 0,
            'image_count': 0
        }
        
        try:
            doc = Document(file_path)
            
            # Extract paragraphs
            full_text = []
            for para in doc.paragraphs:
                if para.text.strip():
                    result['paragraphs'].append(para.text)
                    full_text.append(para.text)
                    
                    # Check for section headers (usually styled)
                    if para.style and para.style.name.startswith('Heading'):
                        result['sections'].append({
                            'type': para.style.name,
                            'title': para.text,
                            'content': []
                        })
                        
            result['text'] = '\n'.join(full_text)
            result['paragraph_count'] = len(result['paragraphs'])
            result['word_count'] = len(result['text'].split())
            
            # Extract tables
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = []
                    for cell in row.cells:
                        row_data.append(cell.text.strip())
                    table_data.append(row_data)
                result['tables'].append(table_data)
                
            result['table_count'] = len(result['tables'])
            
            # Count images (relationships)
            image_count = 0
            for rel in doc.part.rels.values():
                if "image" in rel.reltype:
                    image_count += 1
            result['image_count'] = image_count
            
        except Exception as e:
            self.stats['errors'].append({
                'file': str(file_path),
                'error': f"Text extraction error: {str(e)}"
            })
            
        return result
        
    def classify_document(self, text: str, filename: str, metadata: Dict) -> str:
        """Classify document type based on content"""
        text_lower = text.lower()
        filename_lower = filename.lower()
        
        # Resume/CV
        if any(keyword in text_lower for keyword in ['resume', 'curriculum vitae', 'work experience', 'education', 'skills']):
            return 'resume'
            
        # Letter
        elif any(keyword in text_lower for keyword in ['dear', 'sincerely', 'regards', 'yours truly']):
            return 'letter'
            
        # Contract/Agreement
        elif any(keyword in text_lower for keyword in ['agreement', 'contract', 'terms and conditions', 'party', 'whereas']):
            return 'contract'
            
        # Proposal
        elif any(keyword in text_lower for keyword in ['proposal', 'proposed', 'project description', 'deliverables']):
            return 'proposal'
            
        # Report
        elif any(keyword in text_lower for keyword in ['executive summary', 'findings', 'conclusion', 'analysis']):
            return 'report'
            
        # Minutes/Notes
        elif any(keyword in text_lower for keyword in ['meeting minutes', 'attendees', 'action items']):
            return 'meeting_notes'
            
        else:
            return 'general_document'
            
    def extract_entities(self, text: str) -> List[Dict]:
        """Extract key entities from text"""
        entities = []
        
        # Extract emails
        email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
        emails = re.findall(email_pattern, text)
        for email in set(emails):
            entities.append({
                'type': 'email',
                'value': email
            })
            
        # Extract phone numbers
        phone_pattern = r'[\+]?[(]?[0-9]{3}[)]?[-\s\.]?[(]?[0-9]{3}[)]?[-\s\.]?[0-9]{4,6}'
        phones = re.findall(phone_pattern, text)
        for phone in set(phones):
            if len(phone) >= 10:
                entities.append({
                    'type': 'phone',
                    'value': phone
                })
                
        # Extract URLs
        url_pattern = r'https?://[^\s<>"{}|\\^`\[\]]+'
        urls = re.findall(url_pattern, text)
        for url in set(urls):
            entities.append({
                'type': 'url',
                'value': url
            })
            
        # Extract dates
        date_patterns = [
            r'\d{1,2}[/-]\d{1,2}[/-]\d{2,4}',
            r'\d{4}[/-]\d{1,2}[/-]\d{1,2}',
            r'(?:January|February|March|April|May|June|July|August|September|October|November|December)\s+\d{1,2},?\s+\d{4}'
        ]
        
        for pattern in date_patterns:
            dates = re.findall(pattern, text, re.IGNORECASE)
            for date in set(dates):
                entities.append({
                    'type': 'date',
                    'value': date
                })
                
        return entities
        
    def process_word_file(self, file_path: Path) -> bool:
        """Process a single Word document"""
        try:
            print(f"   Processing: {file_path.name}")
            
            # Get file info
            file_size = file_path.stat().st_size
            file_hash = self.get_file_hash(file_path)
            
            # Extract metadata
            metadata = self.extract_docx_metadata(file_path)
            
            # Extract text and structure
            content = self.extract_text_and_structure(file_path)
            
            # Classify document
            doc_type = self.classify_document(content['text'], file_path.name, metadata)
            
            # Extract entities
            entities = self.extract_entities(content['text'])
            
            # Store in database
            cursor = self.conn.cursor()
            
            cursor.execute('''
                INSERT OR REPLACE INTO word_docs 
                (file_path, file_name, file_size, file_hash,
                 creation_date, modification_date, author, last_modified_by,
                 title, subject, word_count, page_count, paragraph_count,
                 table_count, image_count, document_type, extracted_text,
                 extracted_tables, key_entities)
                VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
            ''', (
                str(file_path),
                file_path.name,
                file_size,
                file_hash,
                metadata.get('created'),
                metadata.get('modified'),
                metadata.get('author'),
                metadata.get('last_modified_by'),
                metadata.get('title'),
                metadata.get('subject'),
                content['word_count'],
                None,  # Page count not easily available
                content['paragraph_count'],
                content['table_count'],
                content['image_count'],
                doc_type,
                content['text'][:10000] if content['text'] else None,
                json.dumps(content['tables']) if content['tables'] else None,
                json.dumps(entities)
            ))
            
            doc_id = cursor.lastrowid
            
            # Store sections
            for i, section in enumerate(content['sections']):
                cursor.execute('''
                    INSERT INTO word_sections (doc_id, section_type, section_title, section_order)
                    VALUES (?, ?, ?, ?)
                ''', (doc_id, section['type'], section['title'], i))
                
            # Store entities
            for entity in entities:
                cursor.execute('''
                    INSERT INTO word_entities (doc_id, entity_type, entity_value)
                    VALUES (?, ?, ?)
                ''', (doc_id, entity['type'], entity['value']))
                
            self.conn.commit()
            
            # Update stats
            self.stats['files_processed'] += 1
            if content['text']:
                self.stats['text_extracted'] += 1
            if metadata.get('author'):
                self.stats['metadata_extracted'] += 1
            self.stats['tables_found'] += content['table_count']
            self.stats['images_found'] += content['image_count']
            
            return True
            
        except Exception as e:
            self.stats['errors'].append({
                'file': str(file_path),
                'error': str(e)
            })
            print(f"      ❌ Error: {str(e)}")
            return False
            
    def generate_summary_report(self):
        """Generate summary report"""
        cursor = self.conn.cursor()
        
        print("\n📊 WORD DOCUMENT PROCESSING REPORT")
        print("=" * 60)
        
        print(f"\n📈 Overall Statistics:")
        print(f"   Total files processed: {self.stats['files_processed']}")
        print(f"   Text extracted: {self.stats['text_extracted']}")
        print(f"   Metadata extracted: {self.stats['metadata_extracted']}")
        print(f"   Tables found: {self.stats['tables_found']}")
        print(f"   Images found: {self.stats['images_found']}")
        
        # Document types
        print(f"\n📄 Document Types:")
        cursor.execute('''
            SELECT document_type, COUNT(*) as count
            FROM word_docs
            GROUP BY document_type
            ORDER BY count DESC
        ''')
        
        for doc_type, count in cursor.fetchall():
            print(f"   {doc_type}: {count} files")
            
        # Authors
        print(f"\n👤 Document Authors:")
        cursor.execute('''
            SELECT author, COUNT(*) as count
            FROM word_docs
            WHERE author IS NOT NULL
            GROUP BY author
            ORDER BY count DESC
        ''')
        
        for author, count in cursor.fetchall():
            print(f"   {author}: {count} documents")
            
        # Word count stats
        cursor.execute('''
            SELECT 
                SUM(word_count) as total_words,
                AVG(word_count) as avg_words,
                MAX(word_count) as max_words
            FROM word_docs
        ''')
        
        stats = cursor.fetchone()
        if stats[0]:
            print(f"\n📝 Word Count Statistics:")
            print(f"   Total words: {stats[0]:,}")
            print(f"   Average per document: {int(stats[1]):,}")
            print(f"   Longest document: {stats[2]:,} words")
            
        # Entities
        print(f"\n🔍 Extracted Entities:")
        cursor.execute('''
            SELECT entity_type, COUNT(*) as count
            FROM word_entities
            GROUP BY entity_type
            ORDER BY count DESC
        ''')
        
        for entity_type, count in cursor.fetchall():
            print(f"   {entity_type}: {count} found")
            
        # Errors
        if self.stats['errors']:
            print(f"\n⚠️ Processing Errors ({len(self.stats['errors'])} files):")
            for err in self.stats['errors']:
                print(f"   - {Path(err['file']).name}: {err['error']}")
                
    def process_all_files(self, directory: str):
        """Process all Word documents in directory"""
        word_files = []
        
        print(f"🔍 Scanning for Word documents in {directory}...")
        
        for root, dirs, files in os.walk(directory):
            for file in files:
                if file.lower().endswith(('.docx', '.doc')) and not file.startswith('~'):
                    word_files.append(Path(root) / file)
                    
        print(f"✅ Found {len(word_files)} Word documents to process\n")
        
        # Process each file
        for file_path in word_files:
            self.process_word_file(file_path)
            
        # Generate report
        self.generate_summary_report()
        
        # Save results
        self.save_detailed_results()
        
    def save_detailed_results(self):
        """Save detailed analysis results"""
        cursor = self.conn.cursor()
        
        print(f"\n📋 Document Details:")
        cursor.execute('''
            SELECT file_name, document_type, author, word_count
            FROM word_docs
            ORDER BY file_name
        ''')
        
        for name, doc_type, author, words in cursor.fetchall():
            author_str = f" by {author}" if author else ""
            print(f"   {name}: {doc_type}{author_str} ({words:,} words)")

def main():
    """Main entry point"""
    print("📝 WORD DOCUMENT PROCESSOR")
    print("=" * 60)
    
    processor = WordProcessor()
    processor.process_all_files("/home/jonclaude/agents/Hinkey/desktop_ingest_job")
    
    print("\n✅ Word document processing complete!")

if __name__ == "__main__":
    main()